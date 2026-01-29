#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Automatic Study Guide Generator
Downloads all course materials up to exam date and creates comprehensive study guide
"""

import json
import requests
from datetime import datetime, timezone
from zoneinfo import ZoneInfo
import os
import sys
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

if sys.platform == "win32":
    import codecs
    sys.stdout = codecs.getwriter("utf-8")(sys.stdout.detach())
    sys.stderr = codecs.getwriter("utf-8")(sys.stderr.detach())

CREDENTIALS_DIR = os.path.join(os.path.expanduser("~"), ".claude", "credentials")

def load_json(filename):
    path = os.path.join(CREDENTIALS_DIR, filename)
    with open(path, 'r') as f:
        return json.load(f)

openai_creds = load_json("openai.json")
OPENAI_API_KEY = openai_creds["credentials"]["api_key"]

def clean_html(text):
    """Remove HTML tags"""
    if not text:
        return ""
    clean = re.sub('<[^<]+?>', '', text)
    clean = re.sub(r'\s+', ' ', clean)
    return clean.strip()

def get_modules_before_exam(canvas_url, course_id, api_token, exam_date):
    """Get all modules published before exam date (ONLY test-relevant content)"""
    headers = {"Authorization": f"Bearer {api_token}"}
    modules_url = f"{canvas_url}/api/v1/courses/{course_id}/modules?include[]=items&per_page=100"

    modules_before_exam = []
    pacific_tz = ZoneInfo("America/Los_Angeles")

    # Module names to exclude (non-content modules)
    exclude_keywords = [
        'course essentials',
        'microsoft teams',
        'library research',
        'library guide',
        'instruction guide',
        'getting started',
        'syllabus',
        'resources',
        'how to',
        'orientation'
    ]

    try:
        response = requests.get(modules_url, headers=headers)
        if response.status_code == 200:
            all_modules = response.json()

            for module in all_modules:
                module_name = module.get('name', '').lower()

                # Skip non-content modules
                if any(keyword in module_name for keyword in exclude_keywords):
                    continue

                # Check if module was published before exam
                unlock_at = module.get('unlock_at')

                if unlock_at:
                    try:
                        # Parse UTC time and convert to Pacific
                        unlock_date_utc = datetime.strptime(unlock_at, "%Y-%m-%dT%H:%M:%SZ")
                        unlock_date_utc = unlock_date_utc.replace(tzinfo=timezone.utc)
                        unlock_date_pacific = unlock_date_utc.astimezone(pacific_tz)

                        # Only include if published before exam
                        if unlock_date_pacific < exam_date:
                            modules_before_exam.append(module)
                    except:
                        # If can't parse date, check module position
                        # Don't auto-include if no valid date
                        pass
                else:
                    # No unlock date - check if it's a week/chapter module
                    if any(indicator in module_name for indicator in ['week', 'chapter', 'ch ', 'module']):
                        modules_before_exam.append(module)

            return modules_before_exam
    except Exception as e:
        print(f"Error fetching modules: {e}")

    return []

def get_page_content(canvas_url, course_id, page_url, api_token):
    """Download full content of a Canvas page"""
    headers = {"Authorization": f"Bearer {api_token}"}

    try:
        # Extract page slug from URL
        page_slug = page_url.split('/')[-1]
        page_api_url = f"{canvas_url}/api/v1/courses/{course_id}/pages/{page_slug}"

        response = requests.get(page_api_url, headers=headers)
        if response.status_code == 200:
            page_data = response.json()
            return {
                'title': page_data.get('title', 'Untitled'),
                'content': clean_html(page_data.get('body', ''))
            }
    except Exception as e:
        print(f"Error fetching page {page_url}: {e}")

    return None

def get_quiz_questions(canvas_url, course_id, api_token, exam_date):
    """Get all quiz questions from student's quiz attempts before exam date"""
    headers = {"Authorization": f"Bearer {api_token}"}
    pacific_tz = ZoneInfo("America/Los_Angeles")

    all_quiz_questions = []

    try:
        # Get all quizzes for the course
        quizzes_url = f"{canvas_url}/api/v1/courses/{course_id}/quizzes?per_page=100"
        response = requests.get(quizzes_url, headers=headers)

        if response.status_code == 200:
            quizzes = response.json()

            for quiz in quizzes:
                # Check if quiz was due/published before exam
                due_at = quiz.get('due_at') or quiz.get('lock_at')

                if due_at:
                    try:
                        due_datetime_utc = datetime.strptime(due_at, "%Y-%m-%dT%H:%M:%SZ")
                        due_datetime_utc = due_datetime_utc.replace(tzinfo=timezone.utc)
                        due_datetime_pacific = due_datetime_utc.astimezone(pacific_tz)

                        # Only include quizzes before the exam
                        if due_datetime_pacific >= exam_date:
                            continue
                    except:
                        pass

                quiz_title = quiz.get('title', 'Quiz')
                quiz_id = quiz.get('id')

                print(f"      - Checking quiz: {quiz_title}")

                # Get student's quiz submissions with attempts
                try:
                    # Get all submissions for this quiz (student's submissions)
                    submissions_url = f"{canvas_url}/api/v1/courses/{course_id}/quizzes/{quiz_id}/submissions"
                    submissions_response = requests.get(submissions_url, headers=headers)

                    if submissions_response.status_code == 200:
                        submissions_data = submissions_response.json()

                        # Handle both array and object responses
                        if isinstance(submissions_data, dict):
                            quiz_submissions = submissions_data.get('quiz_submissions', [])
                        else:
                            quiz_submissions = submissions_data

                        # Look through each submission (attempt)
                        for submission in quiz_submissions:
                            # Get the quiz_submission ID
                            quiz_sub_id = submission.get('id')
                            attempt_num = submission.get('attempt', 1)

                            if quiz_sub_id:
                                # Fetch questions from this specific attempt
                                questions_url = f"{canvas_url}/api/v1/quiz_submissions/{quiz_sub_id}/questions"
                                q_response = requests.get(questions_url, headers=headers)

                                if q_response.status_code == 200:
                                    questions_data = q_response.json()

                                    # Extract questions
                                    questions_list = questions_data.get('quiz_submission_questions', [])

                                    if questions_list:
                                        print(f"         ✅ Found {len(questions_list)} questions from attempt {attempt_num}")

                                    for q in questions_list:
                                        question_text = clean_html(q.get('question_text', ''))
                                        question_type = q.get('question_type', '')

                                        # Get answer choices if available
                                        answers_info = ""
                                        answers = q.get('answers', [])
                                        if answers and question_type in ['multiple_choice_question', 'true_false_question']:
                                            answer_texts = [clean_html(a.get('text', '')) for a in answers if a.get('text')]
                                            if answer_texts:
                                                answers_info = "\n   Options: " + " | ".join(answer_texts[:5])

                                        if question_text:
                                            # Check for duplicates
                                            question_key = question_text[:100]
                                            if not any(q['question'][:100] == question_key for q in all_quiz_questions):
                                                all_quiz_questions.append({
                                                    'quiz': quiz_title,
                                                    'question': question_text + answers_info,
                                                    'type': question_type,
                                                    'attempt': attempt_num
                                                })

                                    # Only need questions from one attempt per quiz
                                    if questions_list:
                                        break

                except Exception as e:
                    print(f"         ⚠️  Could not access: {str(e)[:50]}")

    except Exception as e:
        print(f"   Error fetching quizzes: {e}")

    return all_quiz_questions

def get_all_content_before_exam(canvas_url, course_id, api_token, exam_date):
    """Download ALL course content published before exam"""
    print("📥 Downloading course content...")

    all_content = []

    # Get modules before exam
    modules = get_modules_before_exam(canvas_url, course_id, api_token, exam_date)
    print(f"   Found {len(modules)} modules before exam")

    # Note: Quiz questions are hosted on external platforms (Cengage/McGraw-Hill)
    # and cannot be accessed via Canvas API - AI will generate practice questions instead
    quiz_questions = []

    # Download content from each module
    for module in modules:
        module_name = module.get('name', 'Module')
        print(f"   📚 Processing: {module_name}")

        for item in module.get('items', []):
            item_type = item.get('type')
            item_title = item.get('title', 'Item')

            if item_type == 'Page':
                # Download page content
                page_url = item.get('url') or item.get('html_url', '')
                if page_url:
                    page_content = get_page_content(canvas_url, course_id, page_url, api_token)
                    if page_content and page_content['content']:
                        all_content.append({
                            'type': 'Page',
                            'module': module_name,
                            'title': page_content['title'],
                            'content': page_content['content']
                        })

            elif item_type == 'Assignment':
                # Get assignment description
                if item.get('page_url'):
                    headers = {"Authorization": f"Bearer {api_token}"}
                    try:
                        assignment_url = f"{canvas_url}/api/v1/courses/{course_id}/assignments/{item.get('content_id')}"
                        response = requests.get(assignment_url, headers=headers)
                        if response.status_code == 200:
                            assignment = response.json()
                            description = clean_html(assignment.get('description', ''))
                            if description:
                                all_content.append({
                                    'type': 'Assignment',
                                    'module': module_name,
                                    'title': assignment.get('name', item_title),
                                    'content': description
                                })
                    except:
                        pass

            elif item_type == 'Discussion':
                # Get discussion prompt
                headers = {"Authorization": f"Bearer {api_token}"}
                try:
                    discussion_url = f"{canvas_url}/api/v1/courses/{course_id}/discussion_topics/{item.get('content_id')}"
                    response = requests.get(discussion_url, headers=headers)
                    if response.status_code == 200:
                        discussion = response.json()
                        message = clean_html(discussion.get('message', ''))
                        if message:
                            all_content.append({
                                'type': 'Discussion',
                                'module': module_name,
                                'title': discussion.get('title', item_title),
                                'content': message
                            })
                except:
                    pass

    print(f"   ✅ Downloaded {len(all_content)} content items")
    return all_content

def extract_key_concepts_from_content(content_items, exam_name=""):
    """Use AI to extract key concepts and generate multiple choice questions"""
    print("🤖 Generating study guide with AI...")

    # Extract chapters/weeks from exam name (e.g., "Ch 1, 2, 3, 6, 9, 12")
    import re
    chapters_match = re.search(r'ch[apter]*\s*([\d\s,]+)', exam_name.lower())
    chapters_info = ""
    if chapters_match:
        chapters_info = f"\n\nIMPORTANT: This exam covers ONLY Chapters {chapters_match.group(1).strip()}. Focus your study guide ONLY on these specific chapters."

    # Build content text (process in chunks if needed)
    content_text = ""
    for item in content_items[:25]:  # Limit to prevent token overflow
        content_text += f"\n\n=== {item['type']}: {item['title']} ===\n"
        content_text += item['content'][:2000]  # Limit each item

    prompt = f"""You are analyzing course materials to create a comprehensive study guide for an upcoming exam.

Exam: {exam_name}{chapters_info}

Below are course materials (lectures, pages, assignments, discussions) from ONLY the modules/weeks covered on this exam:

{content_text[:12000]}

Create a comprehensive study guide with:

1. KEY CONCEPTS (organized by topic)
   - List ALL important concepts
   - Brief explanation of each
   - Why it's important

2. KEY IDEAS & THEORIES
   - Main theories/frameworks
   - How they connect
   - Real-world applications

3. MULTIPLE CHOICE PRACTICE QUESTIONS (70-100 questions)
   CRITICAL: ONLY create questions about material that will be on THIS exam (covered chapters/weeks only).
   Do NOT include questions on future topics or material not yet covered.

   IMPORTANT FORMAT:
   - Each question must be multiple choice with 4 answer options (A, B, C, D)
   - Label questions as Q1, Q2, Q3, etc.
   - Format exactly like this:

   Q1. [Question text]
   A) [Option A]
   B) [Option B]
   C) [Option C]
   D) [Option D]

   - Cover definitions, applications, comparisons, analysis
   - Mix difficulty levels (easy, medium, hard)
   - Cover ALL topics from the PROVIDED materials ONLY
   - DO NOT include answers here - answers go in answer key

4. ANSWER KEY
   - List ONLY the correct letter for each question
   - Format: Q1: B, Q2: A, Q3: D, etc.
   - Put this section at the very end

Format clearly with headers. Be comprehensive - this is for exam prep."""

    try:
        headers = {
            "Authorization": f"Bearer {OPENAI_API_KEY}",
            "Content-Type": "application/json"
        }

        payload = {
            "model": "gpt-4o-mini",
            "messages": [
                {"role": "system", "content": "You are an expert tutor creating comprehensive study materials with multiple choice questions."},
                {"role": "user", "content": prompt}
            ],
            "temperature": 0.7,
            "max_tokens": 8000
        }

        response = requests.post(
            "https://api.openai.com/v1/chat/completions",
            headers=headers,
            json=payload,
            timeout=120
        )

        response.raise_for_status()
        result = response.json()
        return result['choices'][0]['message']['content']

    except Exception as e:
        print(f"Error extracting concepts: {e}")
        return None

def create_comprehensive_study_guide(exam_info, content_items, ai_analysis, output_path):
    """Create comprehensive .docx study guide"""

    doc = Document()

    # Title
    title = doc.add_heading(f"Comprehensive Study Guide", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading(f"{exam_info['name']}", level=2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Exam info
    doc.add_paragraph(f"Course: {exam_info['course_name']}")
    doc.add_paragraph(f"Exam Date: {exam_info['due_date']}")
    doc.add_paragraph(f"Points: {exam_info['points']}")
    doc.add_paragraph(f"Materials Covered: {len(content_items)} items")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
    doc.add_paragraph()

    # Exam description
    if exam_info.get('description'):
        doc.add_heading("Exam Description", level=1)
        doc.add_paragraph(clean_html(exam_info['description']))
        doc.add_paragraph()

    # AI-generated comprehensive guide
    doc.add_heading("Study Guide Content", level=1)
    doc.add_paragraph()

    for line in ai_analysis.split('\n'):
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue

        # Detect headers
        if line.startswith('**') and line.endswith('**'):
            heading_text = line.replace('**', '')
            doc.add_heading(heading_text, level=2)
        elif line.startswith('#'):
            heading_text = line.lstrip('#').strip()
            level = min(line.count('#'), 3)
            doc.add_heading(heading_text, level=level)
        elif line.startswith('-') or line.startswith('•') or line.startswith('*'):
            doc.add_paragraph(line[1:].strip(), style='List Bullet')
        elif re.match(r'^\d+\.', line):
            doc.add_paragraph(line, style='List Number')
        else:
            doc.add_paragraph(line)

    doc.add_paragraph()
    doc.add_page_break()

    # All course materials section
    doc.add_heading("Course Materials Reviewed", level=1)
    doc.add_paragraph(f"This study guide was created from {len(content_items)} course materials:")
    doc.add_paragraph()

    current_module = None
    for item in content_items:
        module = item['module']

        if module != current_module:
            doc.add_heading(f"📚 {module}", level=2)
            current_module = module

        # Add item
        item_para = doc.add_paragraph()
        item_para.add_run(f"{item['type']}: ").bold = True
        item_para.add_run(item['title'])

    doc.add_paragraph()

    # Footer
    doc.add_page_break()
    footer = doc.add_paragraph("Generated by Auto Study Guide Generator 🤖")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    # Save
    doc.save(output_path)
    print(f"✅ Comprehensive study guide saved: {output_path}")

def generate_study_guide(exam_assignment, canvas_url, course_id, api_token):
    """Main function to generate comprehensive study guide"""

    print(f"\n🎓 Generating comprehensive study guide for: {exam_assignment['name']}")

    # Extract exam info
    exam_date = exam_assignment.get('due_datetime_parsed', datetime.now())

    exam_info = {
        'name': exam_assignment['name'],
        'course_name': exam_assignment['course_name'],
        'description': exam_assignment.get('description', ''),
        'points': exam_assignment.get('points_possible', 0),
        'due_date': exam_date.strftime('%B %d, %Y at %I:%M %p')
    }

    # Download ALL content before exam
    content_items = get_all_content_before_exam(canvas_url, course_id, api_token, exam_date)

    if not content_items:
        print("⚠️  No content found before exam")
        return None

    # Extract key concepts with AI (pass exam name for chapter filtering)
    ai_analysis = extract_key_concepts_from_content(content_items, exam_assignment['name'])

    if not ai_analysis:
        print("❌ Failed to generate AI analysis")
        return None

    # Create output filename
    safe_name = re.sub(r'[^\w\s-]', '', exam_assignment['name'])
    safe_name = re.sub(r'[-\s]+', '-', safe_name)
    output_dir = os.path.join(os.path.expanduser("~"), "Documents", "Canvas Study Guides")
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, f"{safe_name}-COMPREHENSIVE-study-guide.docx")

    # Create comprehensive study guide
    print("📝 Creating comprehensive study guide...")
    create_comprehensive_study_guide(exam_info, content_items, ai_analysis, output_path)

    return output_path

if __name__ == "__main__":
    print("This script is called by the daily planner")
