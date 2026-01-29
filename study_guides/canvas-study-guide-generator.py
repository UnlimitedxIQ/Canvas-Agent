#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Canvas Study Guide Generator
Interactive tool that creates AI-powered study guides from Canvas resources.
Usage: python canvas-study-guide-generator.py "Midterm Exam"
"""

import json
import requests
from datetime import datetime
import os
import sys
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Fix Windows console encoding
if sys.platform == "win32":
    import codecs
    sys.stdout = codecs.getwriter("utf-8")(sys.stdout.detach())
    sys.stderr = codecs.getwriter("utf-8")(sys.stderr.detach())

# Load credentials
CREDENTIALS_DIR = os.path.join(os.path.expanduser("~"), ".claude", "credentials")

def load_json(filename):
    """Load JSON credentials file"""
    path = os.path.join(CREDENTIALS_DIR, filename)
    with open(path, 'r') as f:
        return json.load(f)

# Load credentials
canvas_creds = load_json("canvas.json")
openai_creds = load_json("openai.json")
telegram_creds = load_json("telegram.json")

CANVAS_INSTANCES = canvas_creds["instances"]
OPENAI_API_KEY = openai_creds["credentials"]["api_key"]
TELEGRAM_BOT_TOKEN = telegram_creds["bots"]["canvas_planner"]["bot_token"]
TELEGRAM_CHAT_ID = telegram_creds["bots"]["canvas_planner"]["chat_id"]

def send_telegram(message):
    """Send Telegram notification"""
    url = f"https://api.telegram.org/bot{TELEGRAM_BOT_TOKEN}/sendMessage"
    requests.post(url, json={
        "chat_id": TELEGRAM_CHAT_ID,
        "text": message,
        "parse_mode": "Markdown"
    })

def clean_html(text):
    """Remove HTML tags"""
    if not text:
        return ""
    clean = re.sub('<[^<]+?>', '', text)
    clean = re.sub(r'\s+', ' ', clean)
    return clean.strip()

def search_canvas_assignment(query):
    """Search for assignment across all Canvas instances"""
    print(f"🔍 Searching for: {query}")

    all_assignments = []

    for instance in CANVAS_INSTANCES:
        headers = {"Authorization": f"Bearer {instance['api_token']}"}

        try:
            # Get courses
            courses_url = f"{instance['url']}/api/v1/courses"
            courses = requests.get(courses_url, headers=headers).json()

            for course in courses:
                if 'id' not in course:
                    continue

                # Get assignments
                assignments_url = f"{instance['url']}/api/v1/courses/{course['id']}/assignments"
                assignments = requests.get(assignments_url, headers=headers)

                if assignments.status_code == 200:
                    for assignment in assignments.json():
                        assignment['course_name'] = course.get('name', 'Unknown')
                        assignment['course_id'] = course['id']
                        assignment['canvas_url'] = instance['url']
                        assignment['api_token'] = instance['api_token']
                        all_assignments.append(assignment)
        except Exception as e:
            print(f"Error searching {instance['name']}: {e}")

    # Search by name
    query_lower = query.lower()
    matches = [a for a in all_assignments if query_lower in a.get('name', '').lower()]

    return matches

def get_assignment_modules(canvas_url, course_id, api_token):
    """Get course modules (which may contain slides, readings, etc.)"""
    headers = {"Authorization": f"Bearer {api_token}"}
    modules_url = f"{canvas_url}/api/v1/courses/{course_id}/modules"

    try:
        response = requests.get(modules_url, headers=headers)
        if response.status_code == 200:
            return response.json()
    except Exception as e:
        print(f"Error fetching modules: {e}")

    return []

def get_module_items(canvas_url, course_id, module_id, api_token):
    """Get items within a module"""
    headers = {"Authorization": f"Bearer {api_token}"}
    items_url = f"{canvas_url}/api/v1/courses/{course_id}/modules/{module_id}/items"

    try:
        response = requests.get(items_url, headers=headers)
        if response.status_code == 200:
            return response.json()
    except Exception as e:
        print(f"Error fetching module items: {e}")

    return []

def generate_study_guide_content(assignment, modules_info):
    """Use GPT-4o-mini to generate study guide content"""

    name = assignment.get('name', 'Assignment')
    course = assignment.get('course_name', 'Course')
    description = clean_html(assignment.get('description', ''))
    points = assignment.get('points_possible', 0)

    # Build context from modules
    resources_text = "\n".join([f"- {item}" for item in modules_info[:20]])  # Limit to first 20

    prompt = f"""You are creating a comprehensive study guide for a college student.

**Assignment:** {name}
**Course:** {course}
**Points:** {points}
**Description:** {description[:1000]}

**Available Course Resources:**
{resources_text if resources_text else "No specific resources listed"}

Create a detailed study guide with:

1. **Key Topics to Master** - List 5-8 main topics/concepts
2. **Study Timeline** - Recommended study schedule (3-7 days)
3. **Focus Areas** - What to prioritize based on likely exam content
4. **Practice Strategies** - How to practice and test yourself
5. **Quick Reference** - Key formulas, definitions, or frameworks
6. **Common Mistakes** - What students typically get wrong

Format your response in clear sections with bullet points. Make it practical and actionable."""

    try:
        headers = {
            "Authorization": f"Bearer {OPENAI_API_KEY}",
            "Content-Type": "application/json"
        }

        payload = {
            "model": "gpt-4o-mini",
            "messages": [
                {"role": "system", "content": "You are an expert academic tutor creating study guides."},
                {"role": "user", "content": prompt}
            ],
            "temperature": 0.7,
            "max_tokens": 2000
        }

        response = requests.post(
            "https://api.openai.com/v1/chat/completions",
            headers=headers,
            json=payload,
            timeout=60
        )

        response.raise_for_status()
        result = response.json()
        return result['choices'][0]['message']['content']

    except Exception as e:
        print(f"Error generating study guide: {e}")
        return None

def create_study_guide_docx(assignment, study_content, modules_info, output_path):
    """Create a formatted .docx study guide"""

    doc = Document()

    # Title
    title = doc.add_heading(f"Study Guide: {assignment.get('name', 'Assignment')}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Course info
    doc.add_paragraph(f"Course: {assignment.get('course_name', 'Unknown')}", style='Heading 2')
    doc.add_paragraph(f"Points: {assignment.get('points_possible', 'N/A')}")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
    doc.add_paragraph()

    # Assignment description
    doc.add_heading("Assignment Overview", level=1)
    description = clean_html(assignment.get('description', 'No description available'))
    doc.add_paragraph(description)
    doc.add_paragraph()

    # AI-generated study content
    doc.add_heading("AI Study Guide", level=1)

    # Parse and format the AI content
    for line in study_content.split('\n'):
        line = line.strip()
        if not line:
            continue

        # Check if it's a heading (starts with ** or #)
        if line.startswith('**') and line.endswith('**'):
            heading_text = line.replace('**', '')
            doc.add_heading(heading_text, level=2)
        elif line.startswith('#'):
            heading_text = line.lstrip('#').strip()
            doc.add_heading(heading_text, level=2)
        else:
            doc.add_paragraph(line)

    doc.add_paragraph()

    # Available resources
    if modules_info:
        doc.add_heading("Course Resources", level=1)
        for resource in modules_info[:30]:
            doc.add_paragraph(resource, style='List Bullet')

    doc.add_paragraph()

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph("Generated by AI-Powered Canvas Study Guide Generator 🤖")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    # Save
    doc.save(output_path)
    print(f"✅ Study guide saved: {output_path}")

def main():
    """Main execution"""

    if len(sys.argv) < 2:
        print("Usage: python canvas-study-guide-generator.py \"Assignment Name\"")
        print("Example: python canvas-study-guide-generator.py \"Midterm Exam\"")
        sys.exit(1)

    query = " ".join(sys.argv[1:])

    print("🤖 Canvas Study Guide Generator")
    print("=" * 50)

    # Search for assignment
    matches = search_canvas_assignment(query)

    if not matches:
        print(f"❌ No assignments found matching: {query}")
        send_telegram(f"❌ No assignments found for: {query}")
        sys.exit(1)

    if len(matches) > 1:
        print(f"📋 Found {len(matches)} matches:")
        for idx, match in enumerate(matches, 1):
            print(f"  {idx}. {match.get('name')} ({match.get('course_name')})")
        print("\nUsing the first match...")

    # Use first match
    assignment = matches[0]
    print(f"\n✅ Found: {assignment.get('name')}")
    print(f"   Course: {assignment.get('course_name')}")

    send_telegram(f"🤖 Generating study guide for:\n*{assignment.get('name')}*\n{assignment.get('course_name')}")

    # Get course modules/resources
    print("\n📚 Fetching course resources...")
    modules = get_assignment_modules(
        assignment['canvas_url'],
        assignment['course_id'],
        assignment['api_token']
    )

    modules_info = []
    for module in modules[:10]:  # Limit to first 10 modules
        module_name = module.get('name', 'Unknown Module')
        items = get_module_items(
            assignment['canvas_url'],
            assignment['course_id'],
            module['id'],
            assignment['api_token']
        )

        for item in items:
            item_title = item.get('title', 'Unknown')
            item_type = item.get('type', 'Unknown')
            modules_info.append(f"{module_name} → {item_title} ({item_type})")

    print(f"   Found {len(modules_info)} resources")

    # Generate AI study guide content
    print("\n🤖 Generating AI study guide content...")
    study_content = generate_study_guide_content(assignment, modules_info)

    if not study_content:
        print("❌ Failed to generate study content")
        sys.exit(1)

    # Create output filename
    safe_name = re.sub(r'[^\w\s-]', '', assignment.get('name', 'study-guide'))
    safe_name = re.sub(r'[-\s]+', '-', safe_name)
    output_dir = os.path.join(os.path.expanduser("~"), "Documents", "Canvas Study Guides")
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, f"{safe_name}-study-guide.docx")

    # Create .docx file
    print("\n📝 Creating Word document...")
    create_study_guide_docx(assignment, study_content, modules_info, output_path)

    # Send success notification
    send_telegram(f"""✅ *Study Guide Complete!*

📋 {assignment.get('name')}
🎓 {assignment.get('course_name')}

📄 Saved to:
`{output_path}`

Open the file to view your AI-generated study guide! 🤖""")

    print("\n✨ Study guide generation complete!")
    print(f"📄 File: {output_path}")

if __name__ == "__main__":
    main()
